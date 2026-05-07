"""File processing service for PDF, DOCX, and other document formats."""

import os
import tempfile


class FileProcessor:
    """Process various document formats and extract text content."""

    SUPPORTED_EXTENSIONS = {
        ".pdf": "PDF documents",
        ".docx": "Microsoft Word documents",
        ".txt": "Plain text files",
        ".md": "Markdown files",
        ".csv": "CSV files",
        ".json": "JSON files",
    }

    async def process_file(self, file_path: str, file_name: str) -> dict:
        """Process a file and extract its text content."""
        extension = os.path.splitext(file_name)[1].lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}. Supported: {list(self.SUPPORTED_EXTENSIONS.keys())}")

        processor = self._get_processor(extension)
        content = await processor(file_path)

        return {
            "file_name": file_name,
            "extension": extension,
            "content": content,
            "char_count": len(content),
            "word_count": len(content.split()),
        }

    async def process_pdf(self, file_path: str) -> str:
        """Extract text from PDF files."""
        try:
            import pypdf

            text = []

            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)

                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text.append(f"--- Page {page_num + 1} ---\n{page_text}")

            return "\n\n".join(text)

        except ImportError:
            return self._fallback_pdf_extract(file_path)
        except Exception as e:
            raise ValueError(f"Failed to process PDF: {str(e)}")

    async def process_docx(self, file_path: str) -> str:
        """Extract text from DOCX files."""
        try:
            from docx import Document

            doc = Document(file_path)
            paragraphs = []

            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    table_text.append(row_text)
                paragraphs.append("\nTable:\n" + "\n".join(table_text))

            return "\n\n".join(paragraphs)

        except ImportError:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")
        except Exception as e:
            raise ValueError(f"Failed to process DOCX: {str(e)}")

    async def process_txt(self, file_path: str) -> str:
        """Read plain text files."""
        with open(file_path, encoding="utf-8") as f:
            return f.read()

    async def process_md(self, file_path: str) -> str:
        """Read Markdown files."""
        return await self.process_txt(file_path)

    async def process_csv(self, file_path: str) -> str:
        """Process CSV files into readable text."""
        import csv

        with open(file_path, encoding="utf-8") as f:
            reader = csv.reader(f)
            rows = list(reader)

        if not rows:
            return ""

        header = rows[0]
        data = rows[1:]

        result = [f"Columns: {', '.join(header)}", ""]

        for i, row in enumerate(data):
            row_dict = dict(zip(header, row))
            result.append(f"Row {i + 1}:")
            for key, value in row_dict.items():
                result.append(f"  {key}: {value}")
            result.append("")

        return "\n".join(result)

    async def process_json(self, file_path: str) -> str:
        """Process JSON files into readable text."""
        import json

        with open(file_path, encoding="utf-8") as f:
            data = json.load(f)

        return json.dumps(data, indent=2)

    def _get_processor(self, extension: str):
        """Get the appropriate processor for a file extension."""
        processors = {
            ".pdf": self.process_pdf,
            ".docx": self.process_docx,
            ".txt": self.process_txt,
            ".md": self.process_md,
            ".csv": self.process_csv,
            ".json": self.process_json,
        }
        return processors.get(extension)

    def _fallback_pdf_extract(self, file_path: str) -> str:
        """Fallback PDF extraction using pdftotext if available."""
        import subprocess

        try:
            result = subprocess.run(
                ["pdftotext", "-layout", file_path, "-"],
                capture_output=True,
                text=True,
            )

            if result.returncode == 0:
                return result.stdout

            raise ValueError("pdftotext not available")

        except FileNotFoundError:
            raise ImportError("pypdf is required for PDF processing. Install with: pip install pypdf")


class FileUploadHandler:
    """Handle file uploads with validation and temporary storage."""

    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

    def __init__(self, upload_dir: str | None = None):
        self.upload_dir = upload_dir or tempfile.gettempdir()
        self.processor = FileProcessor()

    def validate_file(self, file_name: str, file_size: int) -> dict:
        """Validate a file before processing."""
        extension = os.path.splitext(file_name)[1].lower()

        if extension not in FileProcessor.SUPPORTED_EXTENSIONS:
            return {
                "valid": False,
                "error": f"Unsupported file type: {extension}",
            }

        if file_size > self.MAX_FILE_SIZE:
            return {
                "valid": False,
                "error": f"File too large. Max size: {self.MAX_FILE_SIZE // (1024 * 1024)}MB",
            }

        return {"valid": True}

    async def process_upload(self, file_content: bytes, file_name: str) -> dict:
        """Process an uploaded file."""
        validation = self.validate_file(file_name, len(file_content))

        if not validation["valid"]:
            raise ValueError(validation["error"])

        with tempfile.NamedTemporaryFile(
            delete=False,
            suffix=os.path.splitext(file_name)[1],
            dir=self.upload_dir,
        ) as tmp:
            tmp.write(file_content)
            tmp_path = tmp.name

        try:
            result = await self.processor.process_file(tmp_path, file_name)
            return result
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
